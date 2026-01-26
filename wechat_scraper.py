import requests
import json
import time
import random
import os
import re
import base64
from datetime import datetime
from bs4 import BeautifulSoup
import threading
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from docx.shared import Inches
from io import BytesIO
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

class WeChatScraper:
    def __init__(self, cookie, token):
        self.cookie = cookie
        self.token = token
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Cookie': cookie
        }
        # Initialize driver path once
        self.driver_path = self._get_driver_path()
        self.driver = None
        self.driver_lock = threading.Lock()

    def _get_system_chrome_path(self):
        paths = [
            "/usr/bin/chromium",
            "/usr/bin/chromium-browser",
            "/usr/bin/google-chrome-stable",
            "/usr/bin/google-chrome",
            "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
        ]
        for path in paths:
            if os.path.exists(path):
                return path
        return None

    def _get_driver_path(self):
        # We handle driver path in _get_driver logic now mostly, but kept for compatibility
        return None 

    def _get_driver(self):
        if self.driver is None:
            options = webdriver.ChromeOptions()
            options.add_argument('--headless')
            options.add_argument('--disable-gpu')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            
            binary = self._get_system_chrome_path()
            if binary:
                options.binary_location = binary
            
            try:
                service = None
                if binary and "/usr/bin" in binary and os.path.exists("/usr/bin/chromedriver"):
                    service = Service("/usr/bin/chromedriver")
                
                if not service:       
                    try:
                        driver_path = ChromeDriverManager().install()
                        if "THIRD_PARTY_NOTICES" in driver_path:
                            driver_path = os.path.join(os.path.dirname(driver_path), "chromedriver")
                        os.chmod(driver_path, 0o755)
                        service = Service(driver_path)
                    except:
                        pass
                
                if service:
                    self.driver = webdriver.Chrome(service=service, options=options)
                else:
                    self.driver = webdriver.Chrome(options=options)
                    
            except Exception as e:
                print(f"Error initializing driver: {e}")
        return self.driver

    def log(self, message, callback=None):
        if callback:
            callback(message)
        else:
            print(message)

    def get_fakeid(self, name, callback=None):
        """Search for the official account and get its fakeid."""
        url = "https://mp.weixin.qq.com/cgi-bin/searchbiz"
        params = {
            "action": "search_biz",
            "begin": 0,
            "count": 5,
            "query": name,
            "token": self.token,
            "lang": "zh_CN",
            "f": "json",
            "ajax": 1
        }
        
        try:
            self.log(f"Searching for '{name}'...", callback)
            response = requests.get(url, headers=self.headers, params=params)
            data = response.json()
            
            if data.get('base_resp', {}).get('ret') != 0:
                self.log(f"Error searching for account: {data}", callback)
                return None
                
            for item in data.get('list', []):
                if item['nickname'] == name:
                    return item['fakeid']
            
            self.log(f"Account '{name}' not found.", callback)
            return None
        except Exception as e:
            self.log(f"Exception in get_fakeid: {e}", callback)
            return None

    def get_articles(self, fakeid, callback=None, date_range=None):
        """
        Fetch article list for the given fakeid.
        date_range: tuple (start_date, end_date) objects.
        """
        url = "https://mp.weixin.qq.com/cgi-bin/appmsg"
        articles = []
        page = 0
        MAX_PAGES_ESTIMATE = 40 
        
        should_stop = False
        
        while not should_stop:
            self.log(f"Fetching page {page + 1}...", callback)
            params = {
                "token": self.token,
                "lang": "zh_CN",
                "f": "json",
                "ajax": 1,
                "action": "list_ex",
                "begin": page * 5,
                "count": 5,
                "query": "",
                "fakeid": fakeid,
                "type": 9
            }
            
            try:
                response = requests.get(url, headers=self.headers, params=params)
                data = response.json()
                
                if data.get('base_resp', {}).get('ret') == 200013:
                    self.log("⚠️ Rate limit detected (freq control). Waiting 60 seconds...", callback)
                    time.sleep(60)
                    response = requests.get(url, headers=self.headers, params=params)
                    data = response.json()
                    if data.get('base_resp', {}).get('ret') != 0:
                        self.log("❌ Rate limit persists. Please try again in 1-24 hours.", callback)
                        break

                if data.get('base_resp', {}).get('ret') != 0:
                    self.log(f"Error fetching articles: {data}", callback)
                    break
                    
                msg_list = data.get('app_msg_list', [])
                if not msg_list:
                    self.log("No more articles found.", callback)
                    break
                    
                for msg in msg_list:
                    create_time = datetime.fromtimestamp(msg['create_time'])
                    date_str = create_time.strftime('%Y-%m-%d')
                    
                    # Date Filtering
                    if date_range:
                        start_date, end_date = date_range
                        # Convert to date objects for comparison
                        msg_date = create_time.date()
                        
                        if msg_date < start_date:
                            self.log(f"Reached articles older than {start_date}. Stopping.", callback)
                            should_stop = True
                            break
                        
                        if msg_date > end_date:
                            continue # Skip newer articles
                    
                    article_info = {
                        "title": msg['title'],
                        "link": msg['link'],
                        "create_time": msg['create_time'],
                        "date_str": date_str,
                        "digest": msg['digest']
                    }
                    articles.append(article_info)
                
                if should_stop:
                    break
                    
                page += 1
                
                if page > MAX_PAGES_ESTIMATE:
                    self.log("⚠️ Warning: Reached potential WeChat API limit (approx 200 articles). Older articles may not be accessible via this method.", callback)
                
                time.sleep(random.randint(3, 6)) 
                
            except Exception as e:
                self.log(f"Exception in get_articles: {e}", callback)
                break
                
        return articles

    def _clean_filename(self, title):
        return re.sub(r'[\\/*?:"<>|]', "", title)

    def _convert_html_to_pdf_selenium(self, html_path, pdf_path):
        """Convert HTML file to PDF using Selenium (Print to PDF)."""
        # Lock driver usage to ensure thread safety
        with self.driver_lock:
            driver = self._get_driver()
            if not driver:
                return False
                
            try:
                driver.get(f"file://{os.path.abspath(html_path)}")
                
                print_params = {
                    "landscape": False,
                    "displayHeaderFooter": False,
                    "printBackground": True,
                    "preferCSSPageSize": True,
                }
                
                result = driver.execute_cdp_cmd("Page.printToPDF", print_params)
                
                with open(pdf_path, 'wb') as f:
                    f.write(base64.b64decode(result['data']))
                    
                return True
            except Exception as e:
                print(f"Selenium PDF Error: {e}")
                # If driver crashes, reset it
                self.close_driver()
                return False

    def save_article_content(self, article, base_dir, formats=['html'], callback=None):
        """Download and save the article content in specified formats."""
        title = article['title']
        date_str = article['date_str']
        safe_title = self._clean_filename(title)
        filename_base = f"{date_str}_{safe_title}"
        
        # Fetch content once
        try:
            response = requests.get(article['link'], headers=self.headers)
            response.encoding = 'utf-8'
            
            if response.status_code != 200:
                self.log(f"Failed to download {title}: Status {response.status_code}", callback)
                return False
            
            # Process HTML for local viewing
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 1. Add no-referrer meta tag to bypass anti-hotlinking
            meta_tag = soup.new_tag('meta', attrs={"name": "referrer", "content": "no-referrer"})
            if soup.head:
                soup.head.insert(0, meta_tag)
            else:
                # Create head if missing
                head = soup.new_tag('head')
                head.insert(0, meta_tag)
                soup.insert(0, head)
                
            # 2. Fix lazy-loaded images (Embed as Base64)
            for img in soup.find_all('img'):
                img_url = img.get('data-src') or img.get('src')
                if img_url:
                    try:
                        # Download image
                        img_resp = requests.get(img_url, headers=self.headers, timeout=10)
                        if img_resp.status_code == 200:
                            b64_data = base64.b64encode(img_resp.content).decode('utf-8')
                            
                            # Determine mime type
                            mime_type = "image/jpeg"
                            if "png" in img_url: mime_type = "image/png"
                            elif "gif" in img_url: mime_type = "image/gif"
                            elif "svg" in img_url: mime_type = "image/svg+xml"
                            
                            img['src'] = f"data:{mime_type};base64,{b64_data}"
                            
                            # Remove data-src to prevent lazy loading scripts from messing it up
                            if 'data-src' in img.attrs: del img['data-src']
                    except Exception as e:
                        # Fallback: just use the URL if download fails
                        if 'data-src' in img.attrs:
                            img['src'] = img['data-src']
                        print(f"Failed to embed image: {e}")
                    
            html_content = str(soup)
            
        except Exception as e:
            self.log(f"Error downloading {title}: {e}", callback)
            return False

        success = False

        # 1. HTML (Save if requested or needed for PDF)
        save_dir_html = os.path.join(base_dir, "HTML")
        html_filepath = os.path.join(save_dir_html, f"{filename_base}.html")
        
        # Only write HTML to disk if user wants it OR if we need it for PDF generation
        if 'html' in formats or 'pdf' in formats:
            if not os.path.exists(save_dir_html): os.makedirs(save_dir_html)
            
            if not os.path.exists(html_filepath):
                with open(html_filepath, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                if 'html' in formats: success = True
            elif 'html' in formats:
                self.log(f"Skip HTML (Exists): {filename_base}", callback)
                success = True

        # 2. PDF (Selenium)
        if 'pdf' in formats:
            save_dir_pdf = os.path.join(base_dir, "PDF")
            if not os.path.exists(save_dir_pdf): os.makedirs(save_dir_pdf)
            filepath = os.path.join(save_dir_pdf, f"{filename_base}.pdf")
            
            if not os.path.exists(filepath):
                try:
                    # Ensure HTML exists for PDF generation
                    if not os.path.exists(html_filepath):
                         with open(html_filepath, 'w', encoding='utf-8') as f:
                            f.write(html_content)

                    pdf_success = self._convert_html_to_pdf_selenium(html_filepath, filepath)
                    if pdf_success:
                        success = True
                    else:
                        self.log(f"Error converting to PDF", callback)
                except Exception as e:
                    self.log(f"Error converting to PDF: {e}", callback)
            else:
                self.log(f"Skip PDF (Exists): {filename_base}", callback)

            # Cleanup HTML if not requested
            if 'html' not in formats and os.path.exists(html_filepath):
                try:
                    os.remove(html_filepath)
                    # Try to remove HTML dir if empty
                    if not os.listdir(save_dir_html):
                        os.rmdir(save_dir_html)
                except:
                    pass

        # 4. Word (Robust Text Extraction)
        if 'docx' in formats:
            save_dir_word = os.path.join(base_dir, "Word")
            if not os.path.exists(save_dir_word): os.makedirs(save_dir_word)
            filepath = os.path.join(save_dir_word, f"{filename_base}.docx")
            
            if not os.path.exists(filepath):
                try:
                    doc = Document()
                    doc.add_heading(title, 0)
                    doc.add_paragraph(f"发布日期: {date_str}")
                    
                    soup = BeautifulSoup(html_content, 'html.parser')
                    content_div = soup.find(id="js_content")
                    
                    if content_div:
                        # Extract text paragraphs and headings and images
                        for element in content_div.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'img']):
                            # Handle Images
                            if element.name == 'img':
                                try:
                                    img_url = element.get('data-src') or element.get('src')
                                    if img_url:
                                        # Handle Base64
                                        if img_url.startswith('data:'):
                                            try:
                                                # Format: data:image/jpeg;base64,.....
                                                header, encoded = img_url.split(',', 1)
                                                img_data = base64.b64decode(encoded)
                                                img_stream = BytesIO(img_data)
                                                doc.add_picture(img_stream, width=Inches(5.5))
                                            except Exception as e:
                                                print(f"Error adding base64 image: {e}")
                                        # Handle URL
                                        else:
                                            img_resp = requests.get(img_url, headers=self.headers, timeout=10)
                                            if img_resp.status_code == 200:
                                                img_stream = BytesIO(img_resp.content)
                                                doc.add_picture(img_stream, width=Inches(5.5)) # Fit to page
                                except Exception as e:
                                    print(f"Error adding image: {e}")
                            
                            # Handle Text
                            else:
                                text = element.get_text(strip=True)
                                if text:
                                    if element.name.startswith('h'):
                                        level = int(element.name[1])
                                        doc.add_heading(text, level=level)
                                    else:
                                        doc.add_paragraph(text)
                                        
                                # Check for images inside paragraphs (common in WeChat)
                                for img in element.find_all('img'):
                                    try:
                                        img_url = img.get('data-src') or img.get('src')
                                        if img_url:
                                            if img_url.startswith('data:'):
                                                try:
                                                    header, encoded = img_url.split(',', 1)
                                                    img_data = base64.b64decode(encoded)
                                                    img_stream = BytesIO(img_data)
                                                    doc.add_picture(img_stream, width=Inches(5.5))
                                                except Exception as e:
                                                    print(f"Error adding inline base64 image: {e}")
                                            else:
                                                img_resp = requests.get(img_url, headers=self.headers, timeout=10)
                                                if img_resp.status_code == 200:
                                                    img_stream = BytesIO(img_resp.content)
                                                    doc.add_picture(img_stream, width=Inches(5.5))
                                    except Exception as e:
                                        print(f"Error adding inline image: {e}")
                    else:
                        # Fallback if no js_content
                        doc.add_paragraph("无法解析文章内容结构，仅保存纯文本。")
                        doc.add_paragraph(soup.get_text())

                    doc.save(filepath)
                    success = True
                except Exception as e:
                    self.log(f"Error converting to Word: {e}", callback)
            else:
                self.log(f"Skip Word (Exists): {filename_base}", callback)

        if success:
            self.log(f"Downloaded: {filename_base}", callback)
            
        return success
